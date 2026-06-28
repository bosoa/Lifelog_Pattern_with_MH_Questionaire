"""
마크다운 논문을 Word 문서로 변환
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re


def parse_markdown_to_word(md_file: str, output_file: str):
    """마크다운 파일을 Word 문서로 변환"""

    # Word 문서 생성
    doc = Document()

    # 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 마크다운 파일 읽기
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table = None
    table_headers = []

    for line in lines:
        line = line.rstrip()

        # 빈 줄
        if not line.strip():
            if not in_table:
                doc.add_paragraph()
            continue

        # H1 제목
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # H2 제목
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # H3 제목
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # H4 제목
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 수평선
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 80)

        # 표 시작 감지
        elif '|' in line and not in_table:
            # 표 헤더 파싱
            table_headers = [h.strip() for h in line.split('|')[1:-1]]
            in_table = True
            table = None

        # 표 구분선 (무시)
        elif in_table and re.match(r'\|[\s\-:]+\|', line):
            # 실제 테이블 생성
            if table is None:
                table = doc.add_table(rows=1, cols=len(table_headers))
                table.style = 'Light Grid Accent 1'

                # 헤더 행 채우기
                header_cells = table.rows[0].cells
                for i, header in enumerate(table_headers):
                    header_cells[i].text = header
                    # 헤더 셀 볼드 처리
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        # 표 데이터 행
        elif in_table and '|' in line:
            if table is not None:
                row_data = [d.strip() for d in line.split('|')[1:-1]]
                row_cells = table.add_row().cells
                for i, data in enumerate(row_data):
                    if i < len(row_cells):
                        # 굵은 텍스트 처리
                        if '**' in data:
                            cell_p = row_cells[i].paragraphs[0]
                            parts = data.split('**')
                            for j, part in enumerate(parts):
                                run = cell_p.add_run(part)
                                if j % 2 == 1:  # ** 사이의 텍스트
                                    run.bold = True
                        else:
                            row_cells[i].text = data

        # 표 종료
        elif in_table and '|' not in line:
            in_table = False
            table = None
            table_headers = []
            # 다음 컨텐츠 처리를 위해 continue하지 않음

        # 코드 블록
        if line.startswith('```'):
            continue  # 코드 블록 마커는 무시

        # 불릿 포인트
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')

        # 번호 리스트
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')

        # 일반 텍스트 (표가 아닌 경우에만)
        elif not in_table and not line.startswith('#') and '|' not in line:
            # 볼드 및 이탤릭 처리
            p = doc.add_paragraph()

            # **bold** 처리
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # *italic* 처리
                    italic_parts = re.split(r'(\*.*?\*)', part)
                    for ipart in italic_parts:
                        if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                            run = p.add_run(ipart[1:-1])
                            run.italic = True
                        else:
                            p.add_run(ipart)

    # 페이지 번호 추가 (섹션 설정)
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # 저장
    doc.save(output_file)
    print(f"✅ Word 문서 생성 완료: {output_file}")


def main():
    """메인 실행 함수"""
    md_file = "draft_for_paper/coordinate_transformation_paper.md"
    output_file = "draft_for_paper/coordinate_transformation_paper.docx"

    print(f"\n{'='*60}")
    print(f"📄 마크다운 → Word 변환")
    print(f"{'='*60}\n")

    parse_markdown_to_word(md_file, output_file)

    print(f"\n{'='*60}")
    print(f"✅ 변환 완료")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
